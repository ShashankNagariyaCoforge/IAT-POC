# src/utils.py
"""
Utility functions for:
- Filesystem helpers (ensure dirs, sanitize names)
- Schema I/O (load/save JSON)
- Text normalization and mapped-value → polygon matching
- DOCX support (DOCX→PDF conversion via docx2pdf or LibreOffice; text extraction fallback)
- Rendering (confidence colors, overlay drawing)
- Rasterization (PDF to images, image loader)
- Writers (annotated PNG, combined PDF)

Dependencies expected (see requirements.txt):
  azure-ai-documentintelligence (used elsewhere)
  requests
  PyMuPDF (fitz)
  opencv-python (cv2)
  Pillow (PIL)
  reportlab
  docx2pdf (optional)
  python-docx (for DOCX text extraction)
"""

import os
import re
import json
import subprocess
import shutil
from typing import Dict, Any, List, Tuple, Optional
from datetime import datetime

# Rendering & PDF tools
import fitz  # PyMuPDF
import cv2
import numpy as np
from PIL import Image
from reportlab.pdfgen import canvas

# Root output directory (subfolders are created per input file)
OUTPUT_DIR = "output"

__all__ = [
    # FS / I/O
    "OUTPUT_DIR", "ensure_output_dir", "ensure_dir", "sanitize_filename",
    "load_schema", "save_json",
    # Text / matching
    "normalize_text", "flatten_mapped_fields", "union_polygon",
    "find_field_polygons_on_page", "collect_mapped_polygons",
    # DOCX support
    "convert_docx_to_pdf", "extract_text_from_docx",
    # Rendering / rasterization / writers
    "confidence_color_bgr", "draw_field_overlays_cv2",
    "rasterize_pdf_pages", "rasterize_image",
    "save_images_as_single_pdf", "save_annotated_png",
]


# =========================
# Filesystem / I/O helpers
# =========================

def ensure_output_dir() -> None:
    """Ensure the global output directory exists."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)


def ensure_dir(path: str) -> None:
    """Ensure a folder exists."""
    os.makedirs(path, exist_ok=True)


def sanitize_filename(name: str, max_len: int = 120) -> str:
    """
    Sanitize filename for Windows/Linux:
    - Keep alnum, dash, underscore, dot, space
    - Trim to max_len
    - Remove trailing dots/spaces
    """
    if not name:
        name = f"file_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    sanitized = re.sub(r'[^A-Za-z0-9\-\._ ]+', '_', name)
    sanitized = sanitized.strip(' .')
    if len(sanitized) > max_len:
        base, ext = os.path.splitext(sanitized)
        sanitized = base[:max_len - len(ext)] + ext
    if not sanitized:
        sanitized = f"file_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    return sanitized


def load_schema(path: str) -> Dict[str, Any]:
    """Load and return JSON schema from a file path."""
    if not os.path.exists(path):
        raise FileNotFoundError(f"Schema file not found: {path}")
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def save_json(data: Dict[str, Any], path: str) -> str:
    """Save dict as pretty JSON at the given path."""
    ensure_dir(os.path.dirname(path))
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    return path


# =========================
# Text normalization / match
# =========================

def normalize_text(s: str) -> str:
    """Lowercase, strip, collapse internal whitespace."""
    return " ".join((s or "").strip().lower().split())


def flatten_mapped_fields(data: Any, prefix: str = "") -> List[Tuple[str, Any]]:
    """
    Flatten nested dict/list into list of (field_path, value) for scalar values.
    field_path examples: 'policy.number', 'insured[0].name'
    """
    out: List[Tuple[str, Any]] = []
    if isinstance(data, dict):
        for k, v in data.items():
            new_prefix = f"{prefix}.{k}" if prefix else k
            out.extend(flatten_mapped_fields(v, new_prefix))
    elif isinstance(data, list):
        for i, v in enumerate(data):
            new_prefix = f"{prefix}[{i}]"
            out.extend(flatten_mapped_fields(v, new_prefix))
    else:
        out.append((prefix, data))
    return out


def union_polygon(polys: List[List[float]]) -> Optional[List[float]]:
    """
    Approximate union by returning the bounding rectangle that encloses all polygons.
    Returns rectangle as 8 floats (x1,y1,x2,y2,x3,y3,x4,y4), or None if no valid inputs.
    """
    xs, ys = [], []
    for p in polys:
        if not p or len(p) != 8:
            continue
        xs.extend([p[0], p[2], p[4], p[6]])
        ys.extend([p[1], p[3], p[5], p[7]])
    if not xs or not ys:
        return None
    x_min, y_min, x_max, y_max = min(xs), min(ys), max(xs), max(ys)
    return [x_min, y_min, x_max, y_min, x_max, y_max, x_min, y_max]


def find_field_polygons_on_page(field_value: str, page: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Find polygons for a single mapped field value on a page by matching against lines then words.
    Returns list[ { 'polygon': [8 floats], 'confidence': float, 'source': 'line'|'word' } ]
    """
    out: List[Dict[str, Any]] = []
    if not field_value:
        return out

    target = normalize_text(str(field_value))

    # 1) LINE-level contains match
    for ln in page.get("lines", []):
        text = normalize_text(ln.get("text", ""))
        if not text:
            continue
        if target and target in text and ln.get("polygon"):
            out.append({
                "polygon": ln["polygon"],
                "confidence": float(ln.get("confidence", 1.0)),
                "source": "line"
            })

    # 2) WORD-level accumulation (sliding window) if nothing found at line-level
    if not out:
        words = page.get("words", [])
        if words:
            w_norm = [normalize_text(w.get("text", "")) for w in words]
            max_len = min(10, len(words))
            for i in range(len(words)):
                if not w_norm[i]:
                    continue
                concat = w_norm[i]
                polys = [words[i].get("polygon")]
                confs = [float(words[i].get("confidence", 1.0))]
                if target == concat:
                    poly_union = union_polygon(polys)
                    if poly_union:
                        out.append({
                            "polygon": poly_union,
                            "confidence": float(sum(confs) / len(confs)),
                            "source": "word"
                        })
                    continue
                for j in range(i + 1, min(i + max_len, len(words))):
                    if w_norm[j]:
                        concat = concat + " " + w_norm[j]
                        polys.append(words[j].get("polygon"))
                        confs.append(float(words[j].get("confidence", 1.0)))
                        if target == concat:
                            poly_union = union_polygon(polys)
                            if poly_union:
                                out.append({
                                    "polygon": poly_union,
                                    "confidence": float(sum(confs) / len(confs)),
                                    "source": "word"
                                })
                            break
    return out


def collect_mapped_polygons(pages: List[Dict[str, Any]], mapped: Dict[str, Any]) -> List[List[Dict[str, Any]]]:
    """
    Build per-page overlays for mapped data.
    Returns: overlays_per_page = [
      [ { 'field': 'policy_number', 'polygon': [...], 'confidence': 0.93 }, ... ],
      ...
    ]
    """
    overlays_per_page: List[List[Dict[str, Any]]] = [[] for _ in range(len(pages))]
    pairs = flatten_mapped_fields(mapped)

    for field_path, value in pairs:
        val_norm = normalize_text(str(value)) if value is not None else ""
        if not val_norm:
            continue
        for idx, page in enumerate(pages):
            matches = find_field_polygons_on_page(value, page)
            for m in matches:
                if m.get("polygon"):
                    overlays_per_page[idx].append({
                        "field": field_path,
                        "polygon": m["polygon"],
                        "confidence": m.get("confidence", 1.0),
                        "source": m.get("source", "line")
                    })
    return overlays_per_page


# =========================
# DOCX support
# =========================

def convert_docx_to_pdf(docx_path: str, out_pdf_path: str) -> bool:
    """
    Convert DOCX to PDF.
    Tries:
      1) docx2pdf (requires Microsoft Word on Windows/Mac)
      2) LibreOffice 'soffice --headless' (cross-platform)
    Returns True if successful, False otherwise.
    """
    ensure_dir(os.path.dirname(out_pdf_path))

    # Try docx2pdf (best on Windows with Word installed)
    try:
        from docx2pdf import convert as docx2pdf_convert
        docx2pdf_convert(docx_path, out_pdf_path)
        return os.path.exists(out_pdf_path) and os.path.getsize(out_pdf_path) > 0
    except Exception:
        pass

    # Try LibreOffice
    soffice = shutil.which("soffice")
    if soffice:
        out_dir = os.path.dirname(out_pdf_path) or "."
        cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path]
        try:
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            base = os.path.splitext(os.path.basename(docx_path))[0]
            candidate = os.path.join(out_dir, base + ".pdf")
            if os.path.exists(candidate):
                if os.path.abspath(candidate) != os.path.abspath(out_pdf_path):
                    shutil.move(candidate, out_pdf_path)
                return True
        except Exception:
            pass

    return False


def extract_text_from_docx(docx_path: str) -> str:
    """
    Fallback: extract plain text from DOCX if conversion fails (no overlays possible).
    Includes paragraphs and table cell text.
    """
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError(
            "python-docx is required for DOCX text extraction. Install via 'pip install python-docx'"
        ) from e

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    doc = Document(docx_path)
    parts: List[str] = []

    # Paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)

    return "\n".join(parts)


# =========================
# Rendering / PDF helpers
# =========================

def confidence_color_bgr(conf: float) -> Tuple[int, int, int]:
    """
    BGR color for OpenCV based on OCR confidence thresholds:
      green:  >= 0.90
      yellow: >= 0.75
      orange: >= 0.60
      red:    <  0.60
    """
    if conf >= 0.90:
        return (0, 255, 0)         # green
    if conf >= 0.75:
        return (0, 255, 255)       # yellow
    if conf >= 0.60:
        return (0, 165, 255)       # orange
    return (0, 0, 255)             # red


def draw_field_overlays_cv2(page_img: np.ndarray, page_data: Dict[str, Any], field_overlays: List[Dict[str, Any]]) -> np.ndarray:
    """
    Draw rectangles (polylines) for mapped fields only. Each overlay has field name label.
    Scales DI coordinates to raster image size using page width/height when available.
    """
    img = page_img.copy()
    img_h, img_w = img.shape[:2]
    p_w = page_data.get("width")
    p_h = page_data.get("height")

    # Scale from DI coordinate system to raster size
    scale_x = img_w / p_w if p_w and p_w > 0 else 1.0
    scale_y = img_h / p_h if p_h and p_h > 0 else 1.0

    for ov in field_overlays:
        poly = ov.get("polygon")
        if not poly or len(poly) != 8:
            continue
        conf = float(ov.get("confidence", 1.0))
        color = confidence_color_bgr(conf)
        pts = np.array([
            [poly[0] * scale_x, poly[1] * scale_y],
            [poly[2] * scale_x, poly[3] * scale_y],
            [poly[4] * scale_x, poly[5] * scale_y],
            [poly[6] * scale_x, poly[7] * scale_y]
        ], dtype=np.int32)

        # Draw polygon
        cv2.polylines(img, [pts], isClosed=True, color=color, thickness=2)

        # Optional: semi-transparent fill (uncomment to enable)
        # overlay = img.copy()
        # cv2.fillPoly(overlay, [pts], color)
        # img = cv2.addWeighted(overlay, 0.15, img, 0.85, 0)

        # Label near top-left corner
        label = f"{ov.get('field', '')}"
        x_tl, y_tl = int(pts[:, 0].min()), int(pts[:, 1].min()) - 6
        y_tl = max(12, y_tl)
        cv2.putText(img, label, (x_tl, y_tl), cv2.FONT_HERSHEY_SIMPLEX, 0.5, color, 2, cv2.LINE_AA)

    return img


def rasterize_pdf_pages(pdf_path: str, dpi: int = 200) -> List[np.ndarray]:
    """
    Returns list of OpenCV BGR images, one per page.
    """
    doc = fitz.open(pdf_path)
    images: List[np.ndarray] = []
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)

    for i in range(len(doc)):
        page = doc.load_page(i)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, 3)
        img_bgr = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
        images.append(img_bgr)
    return images


def rasterize_image(image_path: str) -> np.ndarray:
    """
    Load an image (e.g., PNG) as BGR np.ndarray using OpenCV.
    """
    img = cv2.imread(image_path, cv2.IMREAD_COLOR)
    if img is None:
        raise FileNotFoundError(f"Failed to read image: {image_path}")
    return img


def save_images_as_single_pdf(images: List[np.ndarray], out_pdf: str) -> str:
    """
    Writes all images into a single PDF with page size matching each image.
    Uses ReportLab canvas; creates temporary PNGs that are cleaned up.
    """
    ensure_dir(os.path.dirname(out_pdf))
    c = canvas.Canvas(out_pdf)
    tmp_paths: List[str] = []

    try:
        for idx, img_bgr in enumerate(images):
            img_rgb = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2RGB)
            pil_img = Image.fromarray(img_rgb)

            tmp_path = os.path.join(os.path.dirname(out_pdf), f"_tmp_page_{idx}.png")
            pil_img.save(tmp_path, format="PNG")
            tmp_paths.append(tmp_path)

            w, h = pil_img.size
            c.setPageSize((w, h))
            c.drawImage(tmp_path, 0, 0, width=w, height=h)
            c.showPage()
    finally:
        c.save()
        # Clean temp files
        for p in tmp_paths:
            try:
                os.remove(p)
            except Exception:
                pass

    return out_pdf


def save_annotated_png(img_bgr: np.ndarray, out_png: str) -> str:
    """Save a single annotated image as PNG."""
    ensure_dir(os.path.dirname(out_png))
    ok = cv2.imwrite(out_png, img_bgr)
    if not ok:
        raise RuntimeError(f"Failed to write image: {out_png}")
    return out_png
