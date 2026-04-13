"""
UW Worksheet Service.

Generates a structured underwriter worksheet for a case by:
  Cat 1 — reading already-extracted fields from classification_results (instant, no LLM)
  Cat 2 — running targeted LLM prompts on blob-stored document text
  Cat 3 — generating narrative sections with LLM using Secura guidelines

Streams sections progressively as each completes.
"""

import json
import logging
import os
from datetime import datetime
from typing import AsyncIterator, Dict, List, Optional, Any

import httpx

from config import settings
from models.uw_worksheet import UWSection, UWWorksheet

logger = logging.getLogger(__name__)

# ── Ordered sections ────────────────────────────────────────────────────────
SECTION_KEYS = [
    "submission_overview",
    "proposed_program",
    "entity_operations",
    "employment_profile",
    "loss_history",
    "internet_research",
    "uw_opinion",
]

SECTION_TITLES = {
    "submission_overview": "1. Submission Overview",
    "proposed_program":    "2. Proposed Program Terms",
    "entity_operations":   "3. Entity & Operations Profile",
    "employment_profile":  "4. Employment Profile",
    "loss_history":        "5. Loss History & Prior Claims",
    "internet_research":   "6. Internet Research Summary",
    "uw_opinion":          "7. UW Opinion & Recommendation",
}

# ── Helpers ──────────────────────────────────────────────────────────────────

def _na(v: Any) -> str:
    if v is None:
        return "N/A"
    s = str(v).strip()
    if not s or s.lower() in ("null", "na", "n/a", "none", "—"):
        return "N/A"
    return s


def _load_guidelines() -> str:
    path = getattr(settings, "uw_guidelines_path", None) or os.path.join(
        os.path.dirname(__file__), "..", "config", "uw_guidelines.md"
    )
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        return "Standard Secura underwriting guidelines apply."


async def _call_llm(system_prompt: str, user_message: str, max_tokens: int = 1500) -> str:
    """Thin Azure OpenAI call — returns plain text (no JSON mode)."""
    endpoint = settings.azure_openai_endpoint
    api_key = settings.azure_openai_api_key
    deployment = settings.azure_openai_deployment
    api_version = settings.azure_openai_api_version

    if not endpoint or not api_key:
        return "_LLM not configured — placeholder text._"

    url = f"{endpoint}openai/deployments/{deployment}/chat/completions?api-version={api_version}"
    payload = {
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ],
        "max_tokens": max_tokens,
        "temperature": 0.3,
    }
    headers = {"api-key": api_key, "Content-Type": "application/json"}

    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(url, json=payload, headers=headers)
            resp.raise_for_status()
            return resp.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        logger.error(f"[UWWorksheet] LLM call failed: {e}")
        return f"_Section could not be generated: {e}_"


# ── Cat 1 sections (instant, no LLM) ─────────────────────────────────────────

def _build_submission_overview(cls: Dict, case: Dict) -> str:
    kf = cls.get("key_fields") or {}
    cat = cls.get("classification_category", "N/A")
    sub_type_map = {"New": "New Business", "BOR": "Broker of Record Change", "Renewal": "Renewal"}
    sub_type = sub_type_map.get(cat, _na(kf.get("submission_type")))

    lines = [
        f"**Applicant Name:** {_na(kf.get('applicant_name') or kf.get('name'))}",
        f"**Effective Date:** {_na(kf.get('effective_date'))}",
        f"**Submission Type:** {sub_type}",
        f"**Business Segment:** Management Liability",
        f"**Secura Product:** {_na(kf.get('secura_product'))}",
        f"**Category:** {_na(cat)}",
        f"**Agency:** {_na(kf.get('agency') or (kf.get('agent') or {}).get('agencyName'))}",
        f"**Agent / Producer:** {_na(kf.get('licensed_producer') or (kf.get('agent') or {}).get('name'))}",
        f"**Agent Email:** {_na(kf.get('agent_email') or (kf.get('agent') or {}).get('email'))}",
        f"**UW / AM:** {_na(kf.get('uw_am'))}",
        f"**Subject Line:** {_na(case.get('subject'))}",
        f"**Received At:** {_na(case.get('created_at', '')[:10])}",
    ]
    return "\n".join(lines)


def _build_proposed_program(cls: Dict) -> str:
    kf = cls.get("key_fields") or {}
    coverages = kf.get("coverages") or []
    exposures = kf.get("exposures") or []

    parts = []

    # Coverage terms
    parts.append("**Coverage Terms**")
    parts.append(f"- Limit of Liability: {_na(kf.get('limit_of_liability'))}")
    parts.append(f"- Deductible: {_na(kf.get('deductible'))}")
    parts.append(f"- Class / Mass Action Ded. Retention: {_na(kf.get('class_mass_action_deductible_retention'))}")
    parts.append(f"- Pending / Prior Litigation Date: {_na(kf.get('pending_or_prior_litigation_date'))}")
    parts.append(f"- Duty to Defend Limit: {_na(kf.get('duty_to_defend_limit'))}")
    parts.append(f"- Defense Outside Limit: {_na(kf.get('defense_outside_limit'))}")

    # Coverages table
    if coverages:
        parts.append("\n**Requested Coverages**")
        for i, c in enumerate(coverages, 1):
            cov = _na(c.get("coverage"))
            lim = _na(c.get("limit"))
            ded = _na(c.get("deductible"))
            desc = _na(c.get("coverageDescription"))
            parts.append(f"{i}. {cov} — Limit: {lim} | Deductible: {ded}")
            if desc and desc != "N/A":
                parts.append(f"   _{desc}_")

    # Exposures
    if exposures:
        parts.append("\n**Reported Exposures**")
        for e in exposures:
            etype = _na(e.get("exposureType"))
            val = _na(e.get("value"))
            desc = _na(e.get("exposureDescription"))
            parts.append(f"- {etype}: {val}" + (f" — {desc}" if desc != "N/A" else ""))

    if not coverages and not exposures:
        parts.append("\n_No specific coverages or exposures extracted from submission._")

    return "\n".join(parts)


# ── Cat 2 sections (targeted LLM on doc text) ────────────────────────────────

async def _get_all_doc_text(docs: List[Dict]) -> str:
    """Read all extracted text blobs for a case and concatenate."""
    from services.blob_storage import BlobStorageService

    blob = BlobStorageService()
    all_text_parts = []

    for doc in docs:
        filename = doc.get("file_name") or doc.get("filename") or "unknown"
        text_path = doc.get("extracted_text_blob_path")
        if not text_path:
            continue
        try:
            container, blob_name = text_path.split("/", 1)
            text = await blob.download_text(container, blob_name)
            if text:
                all_text_parts.append(f"=== DOCUMENT: {filename} ===\n{text}")
        except Exception as e:
            logger.warning(f"[UWWorksheet] Could not read blob {text_path}: {e}")

    return "\n\n".join(all_text_parts) if all_text_parts else ""


async def _build_entity_operations(cls: Dict, doc_text: str) -> str:
    """Cat 1 base + Cat 2 targeted extraction for entity/operations fields."""
    kf = cls.get("key_fields") or {}

    # Cat 1 fields already available
    cat1_fields = {
        "Entity Type": _na(kf.get("entity_type")),
        "Entity Structure": _na(kf.get("entity_structure")),
        "Years in Business": _na(kf.get("years_in_business")),
        "Number of Employees": _na(kf.get("number_of_employees")),
        "Territory Code / States": _na(kf.get("territory_code")),
        "NAICS Code": _na(kf.get("naics_code")),
        "SIC Code": _na(kf.get("sic_code")),
        "Business Description": _na(kf.get("business_description")),
        "Primary Rating State": _na(kf.get("primary_rating_state")),
        "Insured Address": _na(kf.get("address") or (kf.get("insured") or {}).get("address")),
    }

    # Build base from cat 1
    lines = ["**Extracted from Submission Documents**"]
    for k, v in cat1_fields.items():
        lines.append(f"- {k}: {v}")

    # If we have doc text, ask LLM to supplement with any missing fields
    missing = [k for k, v in cat1_fields.items() if v == "N/A"]
    if doc_text and missing:
        sys_prompt = (
            "You are an insurance underwriter assistant. Extract missing entity and operations information "
            "from the provided document text. Return ONLY the requested fields as a brief bulleted list. "
            "If a field is not found, write 'Not found'. Do not fabricate data."
        )
        user_msg = (
            f"From the documents below, extract ONLY these missing fields:\n"
            + "\n".join(f"- {f}" for f in missing)
            + "\n\nDocument text (truncated to 8000 chars):\n"
            + doc_text[:8000]
        )
        supplement = await _call_llm(sys_prompt, user_msg, max_tokens=600)
        lines.append("\n**Additional Information from Documents**")
        lines.append(supplement)

    return "\n".join(lines)


async def _build_employment_profile(cls: Dict, doc_text: str) -> str:
    kf = cls.get("key_fields") or {}

    cat1_fields = {
        "Employment Category": _na(kf.get("employment_category")),
        "EC Number of Employees": _na(kf.get("ec_number_of_employees")),
        "Employee Compensation": _na(kf.get("employee_compensation")),
        "Employees per Band": _na(kf.get("number_of_employees_in_each_band")),
        "Employee Location(s)": _na(kf.get("employee_location")),
        "Employees per Location": _na(kf.get("number_of_employees_in_each_location")),
    }

    lines = ["**Employment Data from Submission**"]
    all_na = True
    for k, v in cat1_fields.items():
        lines.append(f"- {k}: {v}")
        if v != "N/A":
            all_na = False

    if all_na and doc_text:
        sys_prompt = (
            "You are an insurance underwriter assistant. Extract employment and workforce information from "
            "the provided document text. Include: total employees, employment categories (FT/PT), "
            "compensation data, employee locations, and any compensation bands if available. "
            "Return a concise bulleted summary. If no employment data is found, say so."
        )
        user_msg = (
            "Extract employment profile information from these documents "
            "(truncated to 6000 chars):\n\n" + doc_text[:6000]
        )
        result = await _call_llm(sys_prompt, user_msg, max_tokens=600)
        lines = [result]

    return "\n".join(lines)


async def _build_loss_history(doc_text: str) -> str:
    if not doc_text:
        return "_No document text available to extract loss history._"

    sys_prompt = (
        "You are an insurance underwriter assistant. Review the document text and extract any information "
        "about prior claims, loss history, pending litigation, regulatory investigations, or EEOC charges. "
        "Format as a clear structured summary with: (1) Prior Claims, (2) Pending Litigation, "
        "(3) Regulatory Matters, (4) Notes. If no information is found in a category, write 'None identified'. "
        "Do not fabricate data."
    )
    user_msg = (
        "Extract loss history and claims information from these documents "
        "(truncated to 8000 chars):\n\n" + doc_text[:8000]
    )
    return await _call_llm(sys_prompt, user_msg, max_tokens=800)


# ── Cat 3 sections (LLM narrative with guidelines) ───────────────────────────

async def _build_internet_research(enrichment: Optional[Dict]) -> str:
    if not enrichment:
        return "_No web enrichment data available for this case._"

    company_name = enrichment.get("company_name") or "the applicant"
    website = enrichment.get("website")

    lines = [f"**Company:** {company_name}"]
    if website:
        lines.append(f"**Website:** {website}")

    # Pull enriched fields
    enriched_fields = []
    field_map = {
        "entity_type": "Entity Type",
        "entity_structure": "Entity Structure",
        "years_in_business": "Years in Business",
        "number_of_employees": "Number of Employees",
        "territory_code": "Territory / States of Operation",
        "naics_code": "NAICS Code",
    }
    for key, label in field_map.items():
        field_data = enrichment.get(key)
        if field_data and isinstance(field_data, dict):
            val = field_data.get("value")
            conf = field_data.get("confidence", 0)
            if val and val.lower() not in ("null", "n/a", "na", "none"):
                enriched_fields.append(f"- {label}: {val} (confidence: {round(conf * 100)}%)")

    if enriched_fields:
        lines.append("\n**Web-Verified Data Points**")
        lines.extend(enriched_fields)

    # Source URLs
    source_urls = [
        u for u in (enrichment.get("source_urls") or [])
        if u and u.startswith("http") and u not in ("google_search", "source_text")
    ]
    if source_urls:
        lines.append("\n**Research Sources**")
        for url in source_urls[:5]:
            lines.append(f"- {url}")

    # LLM narrative from enriched data
    if enriched_fields:
        sys_prompt = (
            "You are an insurance underwriter. Based on web research data about an applicant, "
            "write a concise 2-3 paragraph narrative summarizing what is publicly known about the company. "
            "Include: business operations, size, any notable public information. "
            "Be factual and professional. Do not speculate beyond what is provided."
        )
        user_msg = (
            f"Company: {company_name}\n"
            + "\n".join(enriched_fields)
            + f"\nSources: {', '.join(source_urls[:3]) if source_urls else 'web research'}"
        )
        narrative = await _call_llm(sys_prompt, user_msg, max_tokens=500)
        lines.append("\n**Research Narrative**")
        lines.append(narrative)

    return "\n".join(lines)


async def _build_uw_opinion(cls: Dict, case: Dict, guidelines: str) -> str:
    kf = cls.get("key_fields") or {}
    cat = cls.get("classification_category", "Unknown")
    summary = cls.get("summary", "")

    context = {
        "Applicant": _na(kf.get("applicant_name") or kf.get("name")),
        "Category": cat,
        "Entity Type": _na(kf.get("entity_type")),
        "Years in Business": _na(kf.get("years_in_business")),
        "Number of Employees": _na(kf.get("number_of_employees")),
        "NAICS Code": _na(kf.get("naics_code")),
        "Territory": _na(kf.get("territory_code")),
        "Business Description": _na(kf.get("business_description")),
        "Limit of Liability": _na(kf.get("limit_of_liability")),
        "Deductible": _na(kf.get("deductible")),
        "Submission Summary": summary,
    }

    context_str = "\n".join(f"- {k}: {v}" for k, v in context.items())

    sys_prompt = (
        "You are a senior underwriter at Secura Insurance specializing in Management Liability. "
        "Based on the submission details and company guidelines provided, write a professional UW opinion. "
        "Structure your response with these sections:\n"
        "**Favorable Factors** — bullet list\n"
        "**Areas of Concern** — bullet list (or 'None identified' if clean)\n"
        "**Subjectivities Required** — bullet list of items needed before binding\n"
        "**Recommendation** — one of: Proceed to Quote / Refer to Senior UW / Decline\n"
        "**Rationale** — 1-2 sentences explaining the recommendation.\n\n"
        "Be professional, specific, and base your opinion on the data provided. "
        "Do not fabricate facts not present in the submission."
    )

    user_msg = (
        "**Submission Details:**\n"
        + context_str
        + "\n\n**Secura Underwriting Guidelines (excerpt):**\n"
        + guidelines[:3000]
    )

    return await _call_llm(sys_prompt, user_msg, max_tokens=900)


# ── Main generator ────────────────────────────────────────────────────────────

async def generate_worksheet_stream(
    case_id: str,
    case: Dict,
    cls: Dict,
    docs: List[Dict],
    enrichment: Optional[Dict],
) -> AsyncIterator[str]:
    """
    Yield SSE-formatted strings for each worksheet section as it completes.

    Format:
        data: {"type": "section_start", "section": "<key>", "title": "<title>"}\\n\\n
        data: {"type": "section_complete", "section": "<key>", "content": "<md>"}\\n\\n
        data: {"type": "done"}\\n\\n
    """

    def _sse(payload: dict) -> str:
        return f"data: {json.dumps(payload)}\n\n"

    guidelines = _load_guidelines()

    # Pre-load all document text once (used for multiple Cat 2 sections)
    doc_text = ""
    try:
        doc_text = await _get_all_doc_text(docs)
    except Exception as e:
        logger.warning(f"[UWWorksheet] Could not load doc text: {e}")

    completed_sections: List[UWSection] = []

    # ── Section 1: Submission Overview (instant) ──────────────────────────
    key = "submission_overview"
    yield _sse({"type": "section_start", "section": key, "title": SECTION_TITLES[key]})
    content = _build_submission_overview(cls, case)
    completed_sections.append(UWSection(section_key=key, title=SECTION_TITLES[key], content=content))
    yield _sse({"type": "section_complete", "section": key, "content": content})

    # ── Section 2: Proposed Program (instant) ────────────────────────────
    key = "proposed_program"
    yield _sse({"type": "section_start", "section": key, "title": SECTION_TITLES[key]})
    content = _build_proposed_program(cls)
    completed_sections.append(UWSection(section_key=key, title=SECTION_TITLES[key], content=content))
    yield _sse({"type": "section_complete", "section": key, "content": content})

    # ── Section 3: Entity & Operations (Cat 2) ───────────────────────────
    key = "entity_operations"
    yield _sse({"type": "section_start", "section": key, "title": SECTION_TITLES[key]})
    content = await _build_entity_operations(cls, doc_text)
    completed_sections.append(UWSection(section_key=key, title=SECTION_TITLES[key], content=content))
    yield _sse({"type": "section_complete", "section": key, "content": content})

    # ── Section 4: Employment Profile (Cat 2) ────────────────────────────
    key = "employment_profile"
    yield _sse({"type": "section_start", "section": key, "title": SECTION_TITLES[key]})
    content = await _build_employment_profile(cls, doc_text)
    completed_sections.append(UWSection(section_key=key, title=SECTION_TITLES[key], content=content))
    yield _sse({"type": "section_complete", "section": key, "content": content})

    # ── Section 5: Loss History (Cat 2) ──────────────────────────────────
    key = "loss_history"
    yield _sse({"type": "section_start", "section": key, "title": SECTION_TITLES[key]})
    content = await _build_loss_history(doc_text)
    completed_sections.append(UWSection(section_key=key, title=SECTION_TITLES[key], content=content))
    yield _sse({"type": "section_complete", "section": key, "content": content})

    # ── Section 6: Internet Research (Cat 3) ─────────────────────────────
    key = "internet_research"
    yield _sse({"type": "section_start", "section": key, "title": SECTION_TITLES[key]})
    content = await _build_internet_research(enrichment)
    completed_sections.append(UWSection(section_key=key, title=SECTION_TITLES[key], content=content))
    yield _sse({"type": "section_complete", "section": key, "content": content})

    # ── Section 7: UW Opinion (Cat 3) ────────────────────────────────────
    key = "uw_opinion"
    yield _sse({"type": "section_start", "section": key, "title": SECTION_TITLES[key]})
    content = await _build_uw_opinion(cls, case, guidelines)
    completed_sections.append(UWSection(section_key=key, title=SECTION_TITLES[key], content=content))
    yield _sse({"type": "section_complete", "section": key, "content": content})

    # ── Persist to DB ─────────────────────────────────────────────────────
    try:
        from services.cosmos_db import CosmosDBService
        db = CosmosDBService()
        worksheet = UWWorksheet(
            case_id=case_id,
            sections=completed_sections,
            generation_status="complete",
        )
        await db.save_uw_worksheet(worksheet)
    except Exception as e:
        logger.error(f"[UWWorksheet] Failed to persist worksheet for {case_id}: {e}")

    yield _sse({"type": "done"})


def build_word_document(worksheet: UWWorksheet, case_id: str) -> bytes:
    """Generate a Word .docx file from the worksheet sections."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = DocxDocument()

    # Title
    title = doc.add_heading("Underwriter Worksheet", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Case ID: {case_id}")
    doc.add_paragraph(f"Generated: {worksheet.generated_at.strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("Secura Insurance — Management Liability")
    doc.add_paragraph("")

    for section in worksheet.sections:
        # Section heading
        doc.add_heading(section.title, level=1)

        # Parse markdown-ish content into paragraphs
        for line in section.content.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
                continue
            # Bold headers (**text**)
            if line.startswith("**") and line.endswith("**") and len(line) > 4:
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            elif line.startswith("- "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("_") and line.endswith("_"):
                p = doc.add_paragraph()
                run = p.add_run(line[1:-1])
                run.italic = True
            else:
                # Inline bold: **key:** value
                if "**" in line:
                    p = doc.add_paragraph()
                    parts = line.split("**")
                    for i, part in enumerate(parts):
                        if not part:
                            continue
                        run = p.add_run(part)
                        run.bold = (i % 2 == 1)
                else:
                    doc.add_paragraph(line)

        doc.add_paragraph("")

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
