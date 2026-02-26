"""
Document parser service (Step 7).
Extracts text from digital PDFs, DOCX files, and images.
Detects if OCR is needed (scanned/handwritten content).
Extracts URLs found within document text.
"""

import io
import logging
import re
from dataclasses import dataclass, field
from typing import List, Optional

import fitz  # PyMuPDF
import docx
from PIL import Image

logger = logging.getLogger(__name__)

URL_PATTERN = re.compile(
    r"https?://[^\s\)\]\>\"\']+",
    re.IGNORECASE,
)

# Min chars threshold to consider text "meaningful" (below this → OCR needed)
OCR_TEXT_THRESHOLD = 50


@dataclass
class ParseResult:
    """Output of the document parser for a single file."""
    raw_text: str = ""
    ocr_required: bool = False
    urls: List[str] = field(default_factory=list)
    file_type: str = ""
    page_count: int = 0


class DocumentParser:
    """Parses PDF, DOCX, and image files to extract text and URLs."""

    async def parse(self, filename: str, content: bytes) -> ParseResult:
        """
        Parse a document and extract text content.

        Args:
            filename: Original filename (used to determine file type).
            content: Raw file bytes.

        Returns:
            ParseResult with extracted text, OCR flag, and URLs.
        """
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        logger.info(f"Parsing document: {filename} (type: {ext})")

        if ext == "pdf":
            return await self._parse_pdf(content)
        elif ext in ("docx", "doc"):
            return await self._parse_docx(content)
        elif ext in ("jpg", "jpeg", "png", "tiff", "tif", "bmp"):
            return await self._parse_image(content, ext)
        else:
            logger.warning(f"Unsupported file type: {ext}. Treating as plain text.")
            try:
                text = content.decode("utf-8", errors="replace")
                return ParseResult(raw_text=text, file_type=ext, urls=self._extract_urls(text))
            except Exception:
                return ParseResult(ocr_required=True, file_type=ext)

    async def _parse_pdf(self, content: bytes) -> ParseResult:
        """
        Extract text from a PDF using PyMuPDF.
        Flags OCR required if extracted text is minimal (scanned/handwritten PDF).
        """
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            pages_text = []
            for page in doc:
                text = page.get_text("text")
                pages_text.append(text)
            full_text = "\n".join(pages_text).strip()
            page_count = len(doc)
            doc.close()

            ocr_required = len(full_text) < OCR_TEXT_THRESHOLD
            urls = self._extract_urls(full_text)

            if ocr_required:
                logger.info(f"PDF appears scanned/handwritten. OCR required. Text length: {len(full_text)}")

            return ParseResult(
                raw_text=full_text,
                ocr_required=ocr_required,
                urls=urls,
                file_type="pdf",
                page_count=page_count,
            )
        except Exception as e:
            logger.error(f"PDF parsing failed: {e}", exc_info=True)
            return ParseResult(ocr_required=True, file_type="pdf")

    async def _parse_docx(self, content: bytes) -> ParseResult:
        """Extract text from a DOCX file using python-docx."""
        try:
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract table text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            full_text = "\n".join(paragraphs)
            urls = self._extract_urls(full_text)
            return ParseResult(
                raw_text=full_text,
                ocr_required=False,
                urls=urls,
                file_type="docx",
            )
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}", exc_info=True)
            return ParseResult(ocr_required=True, file_type="docx")

    async def _parse_image(self, content: bytes, ext: str) -> ParseResult:
        """
        Images always require OCR. Use Pillow to validate the file,
        then flag for OCR processing.
        """
        try:
            image = Image.open(io.BytesIO(content))
            image.verify()
            logger.info(f"Image parsed ({ext}): {image.size if hasattr(image, 'size') else 'unknown'} — OCR required.")
        except Exception as e:
            logger.warning(f"Image validation warning: {e}")
        return ParseResult(ocr_required=True, file_type=ext)

    def _extract_urls(self, text: str) -> List[str]:
        """Extract all URLs from a text string."""
        urls = URL_PATTERN.findall(text)
        # Deduplicate while preserving order
        seen = set()
        unique_urls = []
        for url in urls:
            if url not in seen:
                seen.add(url)
                unique_urls.append(url)
        return unique_urls
